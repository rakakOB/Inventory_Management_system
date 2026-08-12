"""
Generate docs/SRD.docx from SRD.md using python-docx.

Handles: headings (# / ## / ###), markdown tables, inline **bold**,
blockquotes (>), bullet lists (-), and horizontal rules (---).
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
SRC = ROOT / "SRD.md"
OUT_DIR = ROOT / "docs"
OUT = OUT_DIR / "SRD.docx"

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
SEP_RE = re.compile(r"^[\s|:\-]+$")  # markdown table separator row: | --- | :--: |


def add_runs_with_bold(paragraph, text):
    """Add text to a paragraph, converting **bold** markers into bold runs."""
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_horizontal_rule(doc):
    """Insert a thin horizontal line via a paragraph bottom border."""
    p = doc.add_paragraph()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")  # 0.75 pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)


def split_cells(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def parse_table(lines, i):
    """Parse a markdown table starting at lines[i] (header row).

    Returns (rows, next_index); rows[0] is the header.
    """
    header = split_cells(lines[i])
    i += 1
    if i < len(lines) and SEP_RE.match(lines[i].strip()):
        i += 1
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append(split_cells(lines[i]))
        i += 1
    return [header] + rows, i


def add_table(doc, rows):
    """Add a table with a bold header row."""
    header, *body = rows
    table = doc.add_table(rows=1, cols=len(header))
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        table.style = "Table Grid"
    table.autofit = True

    for i, cell_text in enumerate(header):
        p = table.rows[0].cells[i].paragraphs[0]
        add_runs_with_bold(p, cell_text)
        for run in p.runs:
            run.bold = True

    for row in body:
        cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            add_runs_with_bold(cells[i].paragraphs[0], cell_text)

    # slim spacer so the next heading isn't glued to the table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.add_run("").font.size = Pt(4)


def main():
    lines = SRC.read_text(encoding="utf-8").splitlines()

    doc = Document()
    doc.core_properties.title = "System Requirements Document — GIMS v2.0"
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
        elif line == "---":
            add_horizontal_rule(doc)
            i += 1
        elif line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            i += 1
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            i += 1
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=0)  # Title style
            i += 1
        elif line.startswith("> "):
            # merge consecutive blockquote lines into one paragraph
            parts = [line[2:]]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("> "):
                parts.append(lines[i].strip()[2:])
                i += 1
            p = doc.add_paragraph()
            add_runs_with_bold(p, " ".join(parts))
            p.paragraph_format.left_indent = Inches(0.4)
            for run in p.runs:
                run.italic = True
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_bold(p, line[2:])
            i += 1
        else:
            p = doc.add_paragraph()
            add_runs_with_bold(p, line)
            i += 1

    OUT_DIR.mkdir(exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
