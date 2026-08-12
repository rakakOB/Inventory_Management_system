"""Quick verification of docs/SRD.docx contents."""
from docx import Document

doc = Document(r"docs\SRD.docx")
heads = [
    (p.style.name, p.text)
    for p in doc.paragraphs
    if p.style.name.startswith(("Title", "Heading"))
]
print(f"Paragraphs: {len(doc.paragraphs)} | Tables: {len(doc.tables)}")
print("--- Headings ---")
for s, t in heads:
    print(f"[{s}] {t}")
print("--- First table (Master) ---")
for row in doc.tables[0].rows[:3]:
    print(" | ".join(c.text for c in row.cells))
print("--- Last table (Damaged) ---")
for row in doc.tables[-1].rows:
    print(" | ".join(c.text for c in row.cells))
print("--- Blockquote sample ---")
for p in doc.paragraphs:
    if p.runs and p.runs[0].italic:
        print(p.text[:120])
        break
